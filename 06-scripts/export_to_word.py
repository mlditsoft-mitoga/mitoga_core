#!/usr/bin/env python3
"""
Script para convertir documentos Markdown a Word (.docx)
TX Plus - Exportación de Documentación
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import markdown
from bs4 import BeautifulSoup

def setup_styles(doc):
    """Configurar estilos personalizados del documento"""
    styles = doc.styles
    
    # Estilo para código
    try:
        code_style = styles['Code']
    except KeyError:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    
    code_font = code_style.font
    code_font.name = 'Consolas'
    code_font.size = Pt(9)
    code_font.color.rgb = RGBColor(0, 0, 0)
    
    # Fondo gris para código
    pPr = code_style._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)

def add_cover_page(doc, config):
    """Agregar portada profesional"""
    # Título principal
    title = doc.add_heading(config.get('title', 'TX Plus'), 0)
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo
    if config.get('subtitle'):
        subtitle = doc.add_paragraph(config['subtitle'])
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.italic = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Espacio
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Metadata
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if config.get('version'):
        run = meta_para.add_run(f"Versión: {config['version']}\n")
        run.font.size = Pt(11)
    
    if config.get('date'):
        run = meta_para.add_run(f"Fecha: {config['date']}\n")
        run.font.size = Pt(11)
    
    if config.get('audience'):
        run = meta_para.add_run(f"Audiencia: {config['audience']}")
        run.font.size = Pt(11)
    
    # Salto de página
    doc.add_page_break()

def add_header_footer(doc, config):
    """Agregar encabezado y pie de página"""
    section = doc.sections[0]
    
    # Encabezado
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = config.get('header_text', '')
    header_para.style = doc.styles['Header']
    
    # Pie de página con número de página
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Texto del pie
    if config.get('footer_text'):
        footer_para.add_run(f"{config['footer_text']} - Página ")
    else:
        footer_para.add_run("Página ")
    
    # Número de página
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

def parse_markdown_content(md_content):
    """Parsear contenido Markdown y devolver estructura"""
    # Convertir Markdown a HTML
    html = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
    soup = BeautifulSoup(html, 'html.parser')
    
    return soup

def add_table_to_doc(doc, table_html):
    """Agregar tabla HTML al documento Word"""
    rows = table_html.find_all('tr')
    if not rows:
        return
    
    # Contar columnas
    first_row = rows[0]
    cols = len(first_row.find_all(['th', 'td']))
    
    # Crear tabla
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Light Grid Accent 1'
    
    # Llenar celdas
    for i, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for j, cell in enumerate(cells):
            if j < len(table.rows[i].cells):
                cell_text = cell.get_text().strip()
                table.rows[i].cells[j].text = cell_text
                
                # Header en negrita
                if cell.name == 'th':
                    for paragraph in table.rows[i].cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

def add_code_block(doc, code_text):
    """Agregar bloque de código con formato"""
    paragraph = doc.add_paragraph(code_text, style='Code')

def add_diagram_image(doc, diagram_name, image_path):
    """Agregar imagen de diagrama centrada con caption"""
    if os.path.exists(image_path):
        # Agregar imagen centrada
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(6.5))
        
        # Agregar caption
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption.add_run(f"Figura: {diagram_name}")
        caption_run.font.size = Pt(10)
        caption_run.font.italic = True
        caption_run.font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_paragraph()  # Espacio después del diagrama
        print(f"   📊 Imagen agregada: {diagram_name}")
    else:
        print(f"   ⚠️  Imagen no encontrada: {image_path}")

def process_html_element(doc, element, base_path='.'):
    """Procesar elemento HTML y agregarlo al documento"""
    if element.name == 'h1':
        doc.add_heading(element.get_text(), level=1)
    elif element.name == 'h2':
        doc.add_heading(element.get_text(), level=2)
    elif element.name == 'h3':
        doc.add_heading(element.get_text(), level=3)
    elif element.name == 'h4':
        doc.add_heading(element.get_text(), level=4)
    elif element.name == 'p':
        text = element.get_text().strip()
        if text:
            doc.add_paragraph(text)
    elif element.name in ['ul', 'ol']:
        for li in element.find_all('li', recursive=False):
            style = 'List Bullet' if element.name == 'ul' else 'List Number'
            doc.add_paragraph(li.get_text().strip(), style=style)
    elif element.name == 'table':
        add_table_to_doc(doc, element)
    elif element.name in ['pre', 'code']:
        add_code_block(doc, element.get_text())
    elif element.name == 'hr':
        doc.add_paragraph('_' * 80)

def markdown_to_docx(md_file, docx_file, config):
    """
    Convertir archivo Markdown a Word con formato profesional
    
    Args:
        md_file: Path al archivo .md
        docx_file: Path de salida .docx
        config: Dict con configuración
    """
    print(f"📄 Procesando: {md_file}")
    
    # Leer Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Crear documento Word
    doc = Document()
    
    # Configurar estilos
    setup_styles(doc)
    
    # Configurar márgenes (1 pulgada = 914400 EMUs)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Agregar portada si se requiere
    if config.get('add_cover', False):
        add_cover_page(doc, config)
    
    # --- Preprocesar bloques PlantUML en el Markdown: reemplazar fences por marcadores ---
    # Lista de diagramas esperados (orden fijo)
    script_dir = os.path.dirname(os.path.abspath(__file__))
    diagrams_path = os.path.join(script_dir, '04-entregables', 'imagenes', 'diagramas')
    diagrams_files = [
        '01_context.png',
        '02_container.png',
        '03_component.png',
        '04_sequence.png',
        '05_deployment.png',
        '06_state.png'
    ]

    placeholder_map = {}
    def _replace_plantuml(match):
        idx = len(placeholder_map) + 1
        placeholder = f"%%PLANTUML_DIAGRAM_{idx}%%"
        placeholder_map[placeholder] = os.path.join(diagrams_path, diagrams_files[idx-1]) if idx-1 < len(diagrams_files) else None
        return "\n\n" + placeholder + "\n\n"

    # Reemplazar fences ```plantuml ... ``` o ```uml ... ```
    md_content = re.sub(r"```(?:plantuml|uml)[ \t]*\n.*?\n```", _replace_plantuml, md_content, flags=re.DOTALL)

    # Parsear contenido
    soup = parse_markdown_content(md_content)
    
    # Base path para imágenes
    base_path = os.path.dirname(md_file)

    # Procesar elementos: si aparece un placeholder insertar la imagen correspondiente
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol', 'table', 'pre', 'code', 'hr']):
        text = element.get_text()
        # Buscar placeholders generados
        placeholders = [ph for ph in placeholder_map.keys() if ph in text]
        if placeholders:
            # Insertar en el orden en que aparecen en el texto
            for ph in placeholders:
                img_path = placeholder_map.get(ph)
                caption = f"Diagrama: {ph.replace('%%','').replace('_',' ')}"
                if img_path:
                    add_diagram_image(doc, caption, img_path)
                else:
                    print(f"   ⚠️  No hay imagen configurada para {ph}")
            # No insertar el bloque de texto que contenía el PlantUML
            continue

        # Procesamiento normal
        process_html_element(doc, element, base_path)
    
    # Agregar encabezado/pie de página
    if config.get('header_text') or config.get('footer_text'):
        add_header_footer(doc, config)
    
    # Crear directorio de salida si no existe
    os.makedirs(os.path.dirname(docx_file), exist_ok=True)
    
    # Guardar
    doc.save(docx_file)
    print(f"   ✅ Generado: {docx_file}")

def main():
    """Script principal"""
    print("=" * 80)
    print("TX PLUS - EXPORTACIÓN A WORD")
    print("=" * 80)
    print()
    
    # Configuración de documentos a exportar (ORDEN DE LECTURA)
    documents = [
        {
            'name': '01 - Resumen Ejecutivo',
            'md_file': '04-entregables/resumen-ejecutivo.md',
            'docx_file': '04-entregables/word/01_TX_Plus_Resumen_Ejecutivo.docx',
            'config': {
                'title': 'TX Plus - Resumen Ejecutivo',
                'subtitle': 'Rediseño UI/UX de Plataforma TPI (Taxi)',
                'version': '1.0',
                'date': 'Noviembre 2025',
                'audience': 'Stakeholders, Directores',
                'add_cover': True,
                'header_text': 'TX Plus - Resumen Ejecutivo',
                'footer_text': 'Confidencial'
            }
        },
        {
            'name': '02 - Contexto de Negocio',
            'md_file': '01-context-consolidado/01-contexto-negocio.md',
            'docx_file': '04-entregables/word/02_TX_Plus_Contexto_Negocio.docx',
            'config': {
                'title': 'TX Plus - Contexto de Negocio',
                'version': '1.0',
                'date': 'Noviembre 2025',
                'add_cover': True,
                'header_text': 'TX Plus - Contexto de Negocio',
                'footer_text': 'Confidencial'
            }
        },
        {
            'name': '03 - Requisitos Funcionales',
            'md_file': '01-context-consolidado/02-requisitos-funcionales.md',
            'docx_file': '04-entregables/word/03_TX_Plus_Requisitos_Funcionales.docx',
            'config': {
                'title': 'TX Plus - Requisitos Funcionales',
                'version': '1.0',
                'date': 'Noviembre 2025',
                'add_cover': True,
                'header_text': 'TX Plus - Requisitos Funcionales',
                'footer_text': 'Confidencial'
            }
        },
        {
            'name': '04 - Requisitos No Funcionales',
            'md_file': '01-context-consolidado/03-requisitos-no-funcionales.md',
            'docx_file': '04-entregables/word/04_TX_Plus_Requisitos_No_Funcionales.docx',
            'config': {
                'title': 'TX Plus - Requisitos No Funcionales',
                'version': '1.0',
                'date': 'Noviembre 2025',
                'add_cover': True,
                'header_text': 'TX Plus - Requisitos No Funcionales',
                'footer_text': 'Confidencial'
            }
        },
        {
            'name': '05 - Supuestos y Pendientes',
            'md_file': '01-context-consolidado/00-supuestos-y-pendientes.md',
            'docx_file': '04-entregables/word/05_TX_Plus_Supuestos_Pendientes.docx',
            'config': {
                'title': 'TX Plus - Supuestos y Pendientes',
                'version': '1.0',
                'date': 'Noviembre 2025',
                'add_cover': True,
                'header_text': 'TX Plus - Supuestos y Pendientes',
                'footer_text': 'Confidencial'
            }
        },
        {
            'name': '06 - ADR-001: Patrón Arquitectónico',
            'md_file': '03-arquitectura/adrs/001-patron-arquitectonico.md',
            'docx_file': '04-entregables/word/06_TX_Plus_ADR_001_Patron_Arquitectonico.docx',
            'config': {
                'title': 'ADR-001: Patrón Arquitectónico',
                'subtitle': 'Modular Monolith con DDD',
                'version': '1.0',
                'date': 'Noviembre 2025',
                'add_cover': True,
                'header_text': 'TX Plus - ADR-001',
                'footer_text': 'Confidencial'
            }
        },
        {
            'name': '07 - ADR-002: Stack Backend',
            'md_file': '03-arquitectura/adrs/002-stack-tecnologico-backend.md',
            'docx_file': '04-entregables/word/07_TX_Plus_ADR_002_Stack_Backend.docx',
            'config': {
                'title': 'ADR-002: Stack Tecnológico Backend',
                'subtitle': 'Node.js + NestJS + Prisma',
                'version': '1.0',
                'date': 'Noviembre 2025',
                'add_cover': True,
                'header_text': 'TX Plus - ADR-002',
                'footer_text': 'Confidencial'
            }
        },
        {
            'name': '08 - ADR-003: Stack Frontend',
            'md_file': '03-arquitectura/adrs/003-stack-tecnologico-frontend.md',
            'docx_file': '04-entregables/word/08_TX_Plus_ADR_003_Stack_Frontend.docx',
            'config': {
                'title': 'ADR-003: Stack Tecnológico Frontend',
                'subtitle': 'React Native + Next.js',
                'version': '1.0',
                'date': 'Noviembre 2025',
                'add_cover': True,
                'header_text': 'TX Plus - ADR-003',
                'footer_text': 'Confidencial'
            }
        },
        {
            'name': '09 - ADR-004: Estrategia de Datos',
            'md_file': '03-arquitectura/adrs/004-estrategia-datos.md',
            'docx_file': '04-entregables/word/09_TX_Plus_ADR_004_Estrategia_Datos.docx',
            'config': {
                'title': 'ADR-004: Estrategia de Datos',
                'subtitle': 'PostgreSQL + Redis',
                'version': '1.0',
                'date': 'Noviembre 2025',
                'add_cover': True,
                'header_text': 'TX Plus - ADR-004',
                'footer_text': 'Confidencial'
            }
        },
        {
            'name': '10 - Diagramas de Arquitectura',
            'md_file': '03-arquitectura/diagramas/diagramas-arquitectura.md',
            'docx_file': '04-entregables/word/10_TX_Plus_Diagramas_Arquitectura.docx',
            'config': {
                'title': 'TX Plus - Diagramas de Arquitectura',
                'subtitle': 'C4 Model + Deployment',
                'version': '1.0',
                'date': 'Noviembre 2025',
                'add_cover': True,
                'header_text': 'TX Plus - Diagramas',
                'footer_text': 'Confidencial'
            }
        },
        {
            'name': '11 - API Specification',
            'md_file': '03-arquitectura/especificaciones/api-specification.md',
            'docx_file': '04-entregables/word/11_TX_Plus_API_Specification.docx',
            'config': {
                'title': 'TX Plus - API Specification',
                'subtitle': 'REST + WebSocket Endpoints',
                'version': '1.0',
                'date': 'Noviembre 2025',
                'add_cover': True,
                'header_text': 'TX Plus - API Spec',
                'footer_text': 'Confidencial'
            }
        },
        {
            'name': '12 - Roadmap e Implementación',
            'md_file': '03-arquitectura/especificaciones/costos-y-roadmap.md',
            'docx_file': '04-entregables/word/12_TX_Plus_Roadmap_Implementacion.docx',
            'config': {
                'title': 'TX Plus - Roadmap e Implementación',
                'subtitle': 'Plan de 22 Semanas',
                'version': '1.0',
                'date': 'Noviembre 2025',
                'add_cover': True,
                'header_text': 'TX Plus - Roadmap',
                'footer_text': 'Confidencial'
            }
        }
    ]
    
    # Verificar archivos fuente
    print("📋 Verificando archivos fuente...")
    missing = []
    for doc in documents:
        if not Path(doc['md_file']).exists():
            missing.append(doc['md_file'])
            print(f"   ❌ No encontrado: {doc['md_file']}")
        else:
            print(f"   ✅ {doc['name']}")
    
    if missing:
        print(f"\n❌ Faltan {len(missing)} archivos. Abortando.")
        return
    
    print()
    print("🔄 Iniciando conversión...")
    print()
    
    # Procesar cada documento
    success_count = 0
    for i, doc in enumerate(documents, 1):
        print(f"[{i}/{len(documents)}] {doc['name']}")
        try:
            markdown_to_docx(doc['md_file'], doc['docx_file'], doc['config'])
            success_count += 1
        except Exception as e:
            print(f"   ❌ Error: {e}")
        print()
    
    print("=" * 80)
    print(f"✅ EXPORTACIÓN COMPLETADA: {success_count}/{len(documents)} documentos")
    print("=" * 80)
    print()
    print(f"📦 Documentos generados en: 04-entregables/word/")
    print(f"   Total: {success_count} archivos .docx")

if __name__ == "__main__":
    main()
